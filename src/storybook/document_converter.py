"""Document conversion utilities for importing and exporting manuscripts."""

import re
from pathlib import Path

try:
    from docx import Document
    from docx.enum.text import WD_ALIGN_PARAGRAPH

    DOCX_AVAILABLE = True
except ImportError:
    DOCX_AVAILABLE = False

try:
    from pypdf import PdfReader

    PDF_AVAILABLE = True
except ImportError:
    try:
        from PyPDF2 import PdfReader

        PDF_AVAILABLE = True
    except ImportError:
        PDF_AVAILABLE = False


class DocumentConverter:
    """Handles conversion between different document formats."""

    @staticmethod
    def import_from_docx(file_path: Path) -> str:
        """Import a manuscript from a DOCX file.

        Args:
            file_path: Path to the DOCX file

        Returns:
            Manuscript content in Markdown format
        """
        if not DOCX_AVAILABLE:
            raise ImportError("python-docx is required for DOCX import")

        doc = Document(file_path)
        markdown_lines = []

        for paragraph in doc.paragraphs:
            text = paragraph.text.strip()
            if not text:
                markdown_lines.append("")
                continue

            # Detect headings by style or formatting
            if paragraph.style.name.startswith("Heading"):
                level = int(paragraph.style.name[-1]) if paragraph.style.name[-1].isdigit() else 1
                markdown_lines.append(f"{'#' * level} {text}")
            elif paragraph.runs and paragraph.runs[0].bold and len(text) < 100:
                # Likely a heading
                markdown_lines.append(f"## {text}")
            else:
                # Regular paragraph
                markdown_lines.append(text)

        return "\n\n".join([line for line in markdown_lines if line != ""])

    @staticmethod
    def import_from_pdf(file_path: Path) -> str:
        """Import a manuscript from a PDF file.

        Args:
            file_path: Path to the PDF file

        Returns:
            Manuscript content in Markdown format
        """
        if not PDF_AVAILABLE:
            raise ImportError("pypdf or PyPDF2 is required for PDF import")

        reader = PdfReader(file_path)
        text_lines = []

        for page in reader.pages:
            text = page.extract_text()
            text_lines.append(text)

        content = "\n\n".join(text_lines)

        # Clean up the text
        content = re.sub(r"\n{3,}", "\n\n", content)  # Remove excessive newlines
        content = re.sub(r" {2,}", " ", content)  # Remove excessive spaces

        return content

    @staticmethod
    def import_from_text(file_path: Path) -> str:
        """Import a manuscript from a plain text file.

        Args:
            file_path: Path to the text file

        Returns:
            Manuscript content
        """
        return file_path.read_text(encoding="utf-8")

    @staticmethod
    def export_to_docx(content: str, file_path: Path, title: str = "") -> None:
        """Export manuscript to DOCX format.

        Args:
            content: Markdown content
            file_path: Destination path
            title: Manuscript title
        """
        if not DOCX_AVAILABLE:
            raise ImportError("python-docx is required for DOCX export")

        doc = Document()

        # Add title if provided
        if title:
            title_para = doc.add_heading(title, 0)
            title_para.alignment = WD_ALIGN_PARAGRAPH.CENTER

        # Process markdown content
        lines = content.split("\n")
        i = 0

        while i < len(lines):
            line = lines[i].strip()

            if not line:
                i += 1
                continue

            # Headings
            if line.startswith("#"):
                level = len(line) - len(line.lstrip("#"))
                text = line.lstrip("#").strip()
                doc.add_heading(text, min(level, 9))
            else:
                # Regular paragraph
                doc.add_paragraph(line)

            i += 1

        # Save the document
        doc.save(file_path)

    @staticmethod
    def export_to_markdown(content: str, file_path: Path) -> None:
        """Export manuscript to Markdown format.

        Args:
            content: Manuscript content
            file_path: Destination path
        """
        file_path.write_text(content, encoding="utf-8")

    @staticmethod
    def detect_format(file_path: Path) -> str:
        """Detect the format of a file.

        Args:
            file_path: Path to the file

        Returns:
            Format type: 'docx', 'pdf', 'txt', or 'unknown'
        """
        suffix = file_path.suffix.lower()

        if suffix == ".docx":
            return "docx"
        elif suffix == ".pdf":
            return "pdf"
        elif suffix in [".txt", ".md", ".markdown"]:
            return "txt"
        else:
            return "unknown"

    @classmethod
    def import_document(cls, file_path: Path) -> str:
        """Import a document from any supported format.

        Args:
            file_path: Path to the document

        Returns:
            Manuscript content in Markdown format

        Raises:
            ValueError: If format is not supported
        """
        format_type = cls.detect_format(file_path)

        if format_type == "docx":
            return cls.import_from_docx(file_path)
        elif format_type == "pdf":
            return cls.import_from_pdf(file_path)
        elif format_type == "txt":
            return cls.import_from_text(file_path)
        else:
            raise ValueError(f"Unsupported file format: {file_path.suffix}")

    @classmethod
    def export_document(cls, content: str, file_path: Path, title: str = "") -> None:
        """Export a document to the specified format.

        Args:
            content: Manuscript content
            file_path: Destination path
            title: Optional manuscript title

        Raises:
            ValueError: If format is not supported
        """
        format_type = cls.detect_format(file_path)

        if format_type == "docx":
            cls.export_to_docx(content, file_path, title)
        elif format_type in ["txt", "md"]:
            cls.export_to_markdown(content, file_path)
        else:
            raise ValueError(f"Unsupported export format: {file_path.suffix}")
